import base64
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "tests" / "typesetting" / "samples"

PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+0"
    "FtkAAAAASUVORK5CYII="
)


def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def write_png(path: Path) -> None:
    payload = base64.b64decode(PNG_BASE64)
    path.write_bytes(payload)


def save_doc(doc: Document, name: str) -> None:
    doc.save(OUTPUT_DIR / name)


def make_basic_paragraphs() -> None:
    doc = Document()
    title = doc.add_heading("Basic Paragraphs", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("This paragraph mixes ")
    run = p.add_run("bold")
    run.bold = True
    p.add_run(", ")
    run = p.add_run("italic")
    run.italic = True
    p.add_run(", and ")
    run = p.add_run("underlined")
    run.underline = True
    p.add_run(" text.")

    p2 = doc.add_paragraph("Large text with spacing.")
    p2.paragraph_format.space_after = Pt(12)
    p2.runs[0].font.size = Pt(18)

    p3 = doc.add_paragraph("Right aligned paragraph with smaller font.")
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.runs[0].font.size = Pt(10)

    save_doc(doc, "basic-paragraphs.docx")


def make_lists_and_indent() -> None:
    doc = Document()
    doc.add_heading("Lists and Indentation", level=1)

    doc.add_paragraph("Bullet list:")
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Second bullet", style="List Bullet")

    doc.add_paragraph("Numbered list:")
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Second item", style="List Number")

    nested = doc.add_paragraph("Nested bullet", style="List Bullet 2")
    nested.paragraph_format.left_indent = Inches(0.5)

    indented = doc.add_paragraph("Indented paragraph with spacing.")
    indented.paragraph_format.left_indent = Inches(0.75)
    indented.paragraph_format.space_before = Pt(6)

    save_doc(doc, "lists-and-indent.docx")


def make_table_simple() -> None:
    doc = Document()
    doc.add_heading("Simple Table", level=1)
    table = doc.add_table(rows=3, cols=3)
    for row_index, row in enumerate(table.rows, start=1):
        for col_index, cell in enumerate(row.cells, start=1):
            cell.text = f"R{row_index}C{col_index}"
    save_doc(doc, "table-simple.docx")


def make_table_merge() -> None:
    doc = Document()
    doc.add_heading("Merged Cells", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(0, 0).text = "Merged header"
    for row_index in range(1, 3):
        for col_index in range(3):
            table.cell(row_index, col_index).text = f"{row_index}-{col_index}"
    save_doc(doc, "table-merge.docx")


def make_image_inline() -> None:
    doc = Document()
    doc.add_heading("Inline Image", level=1)
    doc.add_paragraph("Image below should scale to width.")
    image_path = OUTPUT_DIR / "sample-image.png"
    write_png(image_path)
    doc.add_picture(str(image_path), width=Inches(3.0))
    doc.add_paragraph("Caption: 3 inch wide image.")
    save_doc(doc, "image-inline.docx")


def make_header_footer() -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Lumina Sample Header"
    section.footer.paragraphs[0].text = "Lumina Sample Footer"
    doc.add_heading("Header and Footer", level=1)
    doc.add_paragraph("This document includes a simple header and footer.")
    doc.add_page_break()
    doc.add_paragraph("Second page to verify header/footer repetition.")
    save_doc(doc, "header-footer.docx")


def make_page_breaks() -> None:
    doc = Document()
    doc.add_heading("Page Breaks", level=1)
    for index in range(1, 4):
        doc.add_paragraph(f"Page {index} content.")
        if index < 3:
            doc.add_page_break()
    save_doc(doc, "page-breaks.docx")


def make_sections_margins() -> None:
    doc = Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Portrait section with default margins.")

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.header.paragraphs[0].text = "Landscape Header"
    doc.add_heading("Section 2", level=1)
    doc.add_paragraph("Landscape section with custom margins.")

    save_doc(doc, "sections-margins.docx")


def make_styles_headings() -> None:
    doc = Document()
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("Body text under heading 1.")
    doc.add_heading("Heading 2", level=2)
    doc.add_paragraph("Body text under heading 2.")
    doc.add_paragraph("Quote style paragraph.", style="Quote")
    save_doc(doc, "styles-headings.docx")


def make_mixed_layout() -> None:
    doc = Document()
    doc.add_heading("Mixed Layout", level=1)
    doc.add_paragraph("Intro paragraph with bold text.").add_run(" Bold!").bold = True

    doc.add_paragraph("Checklist:")
    doc.add_paragraph("Check one", style="List Bullet")
    doc.add_paragraph("Check two", style="List Bullet")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    image_path = OUTPUT_DIR / "sample-image.png"
    if not image_path.exists():
        write_png(image_path)
    doc.add_paragraph("Image: ")
    doc.add_picture(str(image_path), width=Inches(2.5))

    doc.add_page_break()
    doc.add_heading("Second Page", level=1)
    doc.add_paragraph("Content after page break.")
    save_doc(doc, "mixed-layout.docx")


def main() -> None:
    ensure_output_dir()
    make_basic_paragraphs()
    make_lists_and_indent()
    make_table_simple()
    make_table_merge()
    make_image_inline()
    make_header_footer()
    make_page_breaks()
    make_sections_margins()
    make_styles_headings()
    make_mixed_layout()


if __name__ == "__main__":
    main()
